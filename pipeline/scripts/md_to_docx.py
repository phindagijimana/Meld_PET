#!/usr/bin/env python3
"""Convert meldpet.md to meldpet.docx (minimal markdown → Word)."""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt


def add_formatted_run(paragraph, text):
    """Add text with **bold** segments."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            paragraph.add_run(part)


def parse_table(lines):
    rows = []
    for line in lines:
        if re.match(r"^\|[\s\-:|]+\|$", line):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def md_to_docx(md_path, docx_path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = Path(md_path).read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line.strip() == "---":
            i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        if line.startswith("|") and "|" in line[1:]:
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = parse_table(table_lines)
            if rows:
                tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
                tbl.style = "Table Grid"
                for r, row in enumerate(rows):
                    for c, val in enumerate(row):
                        cell = tbl.rows[r].cells[c]
                        cell.text = re.sub(r"\*\*([^*]+)\*\*", r"\1", val)
                        if r == 0:
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.bold = True
            continue

        if line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_run(p, line[2:].strip())
            i += 1
            continue

        if line.strip().startswith("```"):
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        # plain paragraph; strip markdown links to text
        text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", line.strip())
        if text.startswith("*") and text.endswith("*"):
            p = doc.add_paragraph()
            run = p.add_run(text.strip("*"))
            run.italic = True
        else:
            p = doc.add_paragraph()
            add_formatted_run(p, text)
        i += 1

    doc.save(docx_path)
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[2]
    md = root / "meldpet.md"
    out = root / "meldpet.docx"
    if len(sys.argv) > 1:
        md = Path(sys.argv[1])
    if len(sys.argv) > 2:
        out = Path(sys.argv[2])
    md_to_docx(md, out)
